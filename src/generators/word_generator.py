from __future__ import annotations

import base64
import json
import re
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, TypedDict
from dataclasses import dataclass, field

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

REPO_ROOT = Path(__file__).resolve()
REPO_ROOT = REPO_ROOT.parents[2]
OUTPUT_DIR = REPO_ROOT / "generated_documents"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

class WordResult(TypedDict):
    type: str
    filename: str
    mime_type: str
    base64: str
    size_bytes: int
    path: str
    meta: Dict[str, Any]


@dataclass
class WordStyleConfig:
    document: Dict[str, Any] = field(default_factory=dict)
    header: Dict[str, Any] = field(default_factory=dict)
    title: Dict[str, Any] = field(default_factory=dict)
    paragraph: Dict[str, Any] = field(default_factory=dict)
    headings: Dict[str, Any] = field(default_factory=dict)
    table: Dict[str, Any] = field(default_factory=dict)
    lists: Dict[str, Any] = field(default_factory=dict)


def _safe_filename(name: str) -> str:
    s = (name or "document").strip()
    s = re.sub(r"\s+", "_", s)
    s = re.sub(r"[^A-Za-z0-9_\-\.]", "", s)
    return s[:80] if s else "document"


def _file_to_base64(path: Path) -> str:
    return base64.b64encode(path.read_bytes()).decode("utf-8")


def parse_style_config(styling_config: Any) -> WordStyleConfig:
    if not styling_config or styling_config == "{}":
        return WordStyleConfig()
    if isinstance(styling_config, dict):
        return WordStyleConfig(**styling_config)
    return WordStyleConfig(**json.loads(styling_config))


# Style helper functions
def ensure_paragraph_style(doc, name, base="Normal", size=11, bold=False):
    if name in doc.styles:
        return
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles[base]
    style.font.size = Pt(size)
    style.font.bold = bold


def ensure_character_style(doc, name, base=None, bold=False, italic=False):
    if name in doc.styles:
        return
    style = doc.styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
    if base:
        style.base_style = doc.styles[base]
    style.font.bold = bold
    style.font.italic = italic


def initialize_styles(doc):
    ensure_paragraph_style(doc, "Title", size=18, bold=True)
    ensure_paragraph_style(doc, "Body Text", size=11)
    ensure_paragraph_style(doc, "Caption", size=9)

    heading_sizes = {
        1: 16,
        2: 14,
        3: 12,
        4: 11,
        5: 10,
        6: 9,
    }

    for level, size in heading_sizes.items():
        style = doc.styles[f"Heading {level}"]
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(6)

    ensure_character_style(doc, "Emphasis", italic=True)
    ensure_character_style(doc, "Strong", bold=True)


# Header and footer
def add_header_footer(doc, subtitle: str):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = subtitle
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    p = footer.paragraphs[0]
    run = p.add_run()

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')

    instr = OxmlElement('w:instrText')
    instr.text = "PAGE"

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    run._r.extend([fld_begin, instr, fld_end])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Table helper functions
def set_cell_shading(cell, color_hex):
    """Apply background color to a table cell."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    trPr.append(header)


# Internal document generator
def _generate_word_report_internal(
        out_path: Path,
        title: str,
        subtitle: str,
        content: str,
        styling_config: str = "{}",
        page_margins: str = "{}"
) -> Path:
    """
    Internal Word document generation implementation.
    Returns the output file path.
    """
    try:
        content_data = json.loads(content) if isinstance(content, str) else (content or {})
        style_config = parse_style_config(styling_config)
        margins = json.loads(page_margins) if isinstance(page_margins, str) and page_margins else (page_margins or {})

        # Create the document
        doc = Document()
        # Apply styling
        initialize_styles(doc)

        # Page margins
        for section in doc.sections:
            cfg = style_config.document.get("margins", margins)
            section.top_margin = Cm(cfg.get("top", 2))
            section.bottom_margin = Cm(cfg.get("bottom", 2))
            section.left_margin = Cm(cfg.get("left", 2.5))
            section.right_margin = Cm(cfg.get("right", 2.5))

        # HEADER AND FOOTER
        add_header_footer(doc, subtitle)

        # DOCUMENT TITLE
        doc.add_paragraph(title, style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER

        # CONTENT SECTIONS
        for section in content_data.get("sections", []):
            t = section.get("type", "paragraph")

            if t.startswith("heading"):
                try:
                    level = min(int(t.replace("heading", "")), 6)
                except Exception:
                    level = 1
                doc.add_paragraph(section.get("text", ""), style=f"Heading {level}")

            elif t == "paragraph":
                doc.add_paragraph(section["text"], style="Body Text")

            elif t == "table":
                data = section.get("table_data", [])
                if not data or not isinstance(data, list) or not isinstance(data[0], list):
                    doc.add_paragraph("[Invalid table data]", style="Caption")
                else:
                    rows, cols = len(data), len(data[0])
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = style_config.table.get("style", "Light Grid Accent 1")
                    table.autofit = True

                repeat_table_header(table.rows[0])

                for r, row in enumerate(data):
                    row = row[:cols] + [""] * max(0, cols - len(row))
                    for c, value in enumerate(row):
                        cell = table.rows[r].cells[c]
                        cell.text = str(value)
                        if r == 0:
                            set_cell_shading(cell, style_config.table.get("header_shading", "2C5282"))

                if section.get("merge"):
                    for m in section["merge"]:
                        table.cell(*m[:2]).merge(table.cell(*m[2:]))

            elif t == "caption":
                doc.add_paragraph(section["text"], style="Caption")

            elif t == "page_break":
                doc.add_page_break()

            doc.add_paragraph()

        # Save
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return out_path

    except Exception as e:
        raise ValueError(f"Word generation failed: {str(e)}")


def generate_word(
        template_path: Optional[str] = None,
        output_path: Optional[str] = None,
        data: Optional[Dict[str, Any]] = None,
        **kwargs,
) -> WordResult:
    """
    MCP-friendly wrapper for Word document generation.
    Returns base64 content and metadata.
    """
    if not data:
        raise ValueError("Missing 'data' parameter with Word document configuration")

    title = data.get("title", "Document")
    subtitle = data.get("subtitle", "")
    content = data.get("content", {})
    styling_config = data.get("styling_config", {})
    page_margins = data.get("page_margins", {})

    content_str = content if isinstance(content, str) else json.dumps(content)
    styling_str = styling_config if isinstance(styling_config, str) else json.dumps(styling_config)
    margins_str = page_margins if isinstance(page_margins, str) else json.dumps(page_margins)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe = _safe_filename(title)
    filename = f"{safe}_{timestamp}.docx"
    out_path = Path(output_path) if output_path else (OUTPUT_DIR / filename)

    generated_path = _generate_word_report_internal(
        out_path=out_path,
        title=title,
        subtitle=subtitle,
        content=content_str,
        styling_config=styling_str,
        page_margins=margins_str,
    )

    file_bytes = generated_path.read_bytes()
    MAX_BYTES = 8 * 1024 * 1024
    if len(file_bytes) > MAX_BYTES:
        raise RuntimeError(f"Generated DOCX is too large ({len(file_bytes)} bytes). Reduce content or tables.")

    try:
        content_data = json.loads(content_str)
        section_count = len(content_data.get("sections", []))
    except Exception:
        section_count = 0

    return {
        "type": "file_base64",
        "filename": generated_path.name,
        "path": str(generated_path),
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "base64": base64.b64encode(file_bytes).decode("utf-8"),
        "size_bytes": len(file_bytes),
        "meta": {
            "title": title,
            "subtitle": subtitle,
            "sections": section_count,
            "timestamp": timestamp,
        },
    }